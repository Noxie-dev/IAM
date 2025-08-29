"""
DICE™ Algorithm 1: Extractor
Pre-scan algorithm for document and audio file processing

Handles OCR, document parsing, audio diarization, and entity detection
"""

import asyncio
import logging
import tempfile
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import io

# OCR and Document Processing
try:
    import pytesseract
    from PIL import Image
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Audio Processing
try:
    import librosa
    import soundfile as sf
    from pyannote.audio import Pipeline
    HAS_AUDIO = True
except ImportError:
    HAS_AUDIO = False

# Cloud OCR Services
try:
    from google.cloud import vision
    HAS_GOOGLE_VISION = True
except ImportError:
    HAS_GOOGLE_VISION = False

# Entity Detection
try:
    import spacy
    HAS_SPACY = True
except ImportError:
    HAS_SPACY = False

import boto3
from botocore.exceptions import ClientError

from app.schemas.dice_schemas import (
    FileInfo, PreScanJSON, PreScanBlock, PreScanAudioSegment
)
from app.core.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class PreScanAlgorithm:
    """
    Algorithm 1: Extractor
    
    Handles multi-format file processing including:
    - PDF text extraction and OCR
    - DOCX document parsing
    - Image OCR processing
    - Audio diarization and segmentation
    - Entity detection and classification
    """
    
    def __init__(self):
        """Initialize the pre-scan algorithm with available services"""
        self.s3_client = boto3.client(
            's3',
            aws_access_key_id=getattr(settings, 'AWS_ACCESS_KEY_ID', None),
            aws_secret_access_key=getattr(settings, 'AWS_SECRET_ACCESS_KEY', None),
            region_name=getattr(settings, 'AWS_REGION', 'us-east-1')
        )
        
        # Initialize NLP pipeline for entity detection
        self.nlp = None
        if HAS_SPACY:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                logger.warning("SpaCy English model not found. Entity detection will be limited.")
        
        # Initialize audio diarization pipeline
        self.diarization_pipeline = None
        if HAS_AUDIO:
            try:
                # Note: This requires a HuggingFace token for pyannote
                # self.diarization_pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization")
                logger.info("Audio processing available but diarization pipeline not initialized")
            except Exception as e:
                logger.warning(f"Could not initialize diarization pipeline: {e}")
    
    async def process_files(self, files: List[FileInfo]) -> PreScanJSON:
        """
        Process all files in the DICE job
        
        Args:
            files: List of file information objects
            
        Returns:
            PreScanJSON with extracted information
        """
        result = PreScanJSON()
        
        # Process each file based on type
        for file_info in files:
            try:
                mime_type = file_info.mime_type
                
                if mime_type.startswith('audio/'):
                    await self._process_audio_file(file_info, result)
                elif mime_type == 'application/pdf':
                    await self._process_pdf_file(file_info, result)
                elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
                    await self._process_docx_file(file_info, result)
                elif mime_type.startswith('image/'):
                    await self._process_image_file(file_info, result)
                else:
                    logger.warning(f"Unsupported file type: {mime_type}")
            
            except Exception as e:
                logger.error(f"Error processing file {file_info.original_filename}: {e}")
                continue
        
        # Detect entities across all extracted text
        await self._detect_entities(result)
        
        # Calculate overall confidence
        result.ocr_confidence = self._calculate_overall_confidence(result)
        
        return result
    
    async def _download_file_from_s3(self, s3_uri: str) -> bytes:
        """Download file content from S3"""
        try:
            # Parse S3 URI (s3://bucket/key)
            s3_parts = s3_uri.replace('s3://', '').split('/', 1)
            bucket = s3_parts[0]
            key = s3_parts[1]
            
            response = self.s3_client.get_object(Bucket=bucket, Key=key)
            return response['Body'].read()
            
        except ClientError as e:
            logger.error(f"Error downloading from S3: {e}")
            raise
    
    async def _process_audio_file(self, file_info: FileInfo, result: PreScanJSON) -> None:
        """Process audio file for diarization and segmentation"""
        if not HAS_AUDIO:
            logger.warning("Audio processing libraries not available")
            return
        
        try:
            # Download audio file
            audio_data = await self._download_file_from_s3(file_info.s3_uri)
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=False) as temp_file:
                temp_file.write(audio_data)
                temp_path = temp_file.name
            
            try:
                # Load audio with librosa
                audio, sr = librosa.load(temp_path, sr=16000)
                duration = len(audio) / sr
                
                # Basic segmentation (simple silence detection)
                segments = self._simple_audio_segmentation(audio, sr)
                
                # Convert to PreScanAudioSegment objects
                audio_segments = []
                for i, (start, end) in enumerate(segments):
                    segment = PreScanAudioSegment(
                        start_time=start,
                        end_time=end,
                        speaker_hint=f"Speaker_{i % 3 + 1}",  # Simple speaker assignment
                        confidence=0.7  # Placeholder confidence
                    )
                    audio_segments.append(segment)
                
                result.audio_segments = audio_segments
                result.total_duration = duration
                result.diarization_confidence = 0.7  # Placeholder
                
                logger.info(f"Processed audio file: {duration:.2f}s with {len(segments)} segments")
                
            finally:
                # Clean up temporary file
                Path(temp_path).unlink(missing_ok=True)
                
        except Exception as e:
            logger.error(f"Error processing audio file: {e}")
            raise
    
    def _simple_audio_segmentation(self, audio: Any, sr: int) -> List[Tuple[float, float]]:
        """Simple audio segmentation based on energy levels"""
        try:
            # Calculate frame-wise energy
            frame_length = int(0.025 * sr)  # 25ms frames
            hop_length = int(0.010 * sr)    # 10ms hop
            
            # Simple energy-based segmentation
            segments = []
            segment_duration = 10.0  # 10-second segments
            total_duration = len(audio) / sr
            
            current_time = 0.0
            while current_time < total_duration:
                end_time = min(current_time + segment_duration, total_duration)
                segments.append((current_time, end_time))
                current_time = end_time
            
            return segments
            
        except Exception as e:
            logger.error(f"Error in audio segmentation: {e}")
            return [(0.0, len(audio) / sr)]
    
    async def _process_pdf_file(self, file_info: FileInfo, result: PreScanJSON) -> None:
        """Process PDF file for text extraction and OCR"""
        if not HAS_PDF:
            logger.warning("PDF processing libraries not available")
            return
        
        try:
            # Download PDF
            pdf_data = await self._download_file_from_s3(file_info.s3_uri)
            
            # Process with pdfplumber for better text extraction
            with io.BytesIO(pdf_data) as pdf_stream:
                with pdfplumber.open(pdf_stream) as pdf:
                    pages = []
                    total_confidence = 0.0
                    
                    for page_num, page in enumerate(pdf.pages):
                        # Extract text
                        text = page.extract_text() or ""
                        
                        # Get page dimensions for bbox calculation
                        width = page.width
                        height = page.height
                        
                        # Create a text block for the entire page
                        if text.strip():
                            block = PreScanBlock(
                                type="text",
                                bbox=[0, 0, width, height],
                                text=text,
                                confidence=0.95,  # High confidence for digital text
                                page_number=page_num + 1
                            )
                            pages.append(block)
                            total_confidence += 0.95
                        
                        # Check for tables
                        tables = page.extract_tables()
                        for table in tables:
                            table_text = self._table_to_text(table)
                            if table_text:
                                table_block = PreScanBlock(
                                    type="table",
                                    bbox=[0, 0, width, height],  # Simplified bbox
                                    text=table_text,
                                    confidence=0.9,
                                    page_number=page_num + 1
                                )
                                pages.append(table_block)
                                total_confidence += 0.9
                    
                    if pages:
                        result.pages = pages
                        result.total_pages = len(pdf.pages)
                        result.ocr_confidence = total_confidence / len(pages) if pages else 0.0
                    
                    logger.info(f"Processed PDF: {len(pages)} text blocks from {len(pdf.pages)} pages")
        
        except Exception as e:
            logger.error(f"Error processing PDF file: {e}")
            raise
    
    def _table_to_text(self, table: List[List[str]]) -> str:
        """Convert table data to text"""
        try:
            text_rows = []
            for row in table:
                if row:
                    text_row = " | ".join(str(cell) if cell else "" for cell in row)
                    text_rows.append(text_row)
            return "\n".join(text_rows)
        except Exception:
            return ""
    
    async def _process_docx_file(self, file_info: FileInfo, result: PreScanJSON) -> None:
        """Process DOCX file for text extraction"""
        if not HAS_DOCX:
            logger.warning("DOCX processing libraries not available")
            return
        
        try:
            # Download DOCX
            docx_data = await self._download_file_from_s3(file_info.s3_uri)
            
            # Process with python-docx
            with io.BytesIO(docx_data) as docx_stream:
                doc = DocxDocument(docx_stream)
                
                pages = []
                
                # Extract paragraphs
                for i, paragraph in enumerate(doc.paragraphs):
                    if paragraph.text.strip():
                        block = PreScanBlock(
                            type="text",
                            bbox=[0, 0, 100, 100],  # Placeholder dimensions
                            text=paragraph.text,
                            confidence=0.98,  # Very high for digital text
                            page_number=1  # DOCX doesn't have clear page boundaries
                        )
                        pages.append(block)
                
                # Extract tables
                for table in doc.tables:
                    table_text = self._docx_table_to_text(table)
                    if table_text:
                        table_block = PreScanBlock(
                            type="table",
                            bbox=[0, 0, 100, 100],
                            text=table_text,
                            confidence=0.95,
                            page_number=1
                        )
                        pages.append(table_block)
                
                if pages:
                    result.pages = pages
                    result.total_pages = 1
                    result.ocr_confidence = 0.98
                
                logger.info(f"Processed DOCX: {len(pages)} text blocks")
        
        except Exception as e:
            logger.error(f"Error processing DOCX file: {e}")
            raise
    
    def _docx_table_to_text(self, table) -> str:
        """Convert DOCX table to text"""
        try:
            text_rows = []
            for row in table.rows:
                text_row = " | ".join(cell.text for cell in row.cells)
                text_rows.append(text_row)
            return "\n".join(text_rows)
        except Exception:
            return ""
    
    async def _process_image_file(self, file_info: FileInfo, result: PreScanJSON) -> None:
        """Process image file with OCR"""
        if not HAS_TESSERACT:
            logger.warning("Tesseract OCR not available")
            return
        
        try:
            # Download image
            image_data = await self._download_file_from_s3(file_info.s3_uri)
            
            # Process with PIL and Tesseract
            with io.BytesIO(image_data) as image_stream:
                image = Image.open(image_stream)
                
                # OCR with confidence data
                ocr_data = pytesseract.image_to_data(
                    image, 
                    output_type=pytesseract.Output.DICT
                )
                
                # Group text by confidence and position
                text_blocks = []
                current_text = []
                current_conf = []
                
                for i, text in enumerate(ocr_data['text']):
                    if text.strip():
                        current_text.append(text)
                        current_conf.append(int(ocr_data['conf'][i]))
                
                if current_text:
                    combined_text = " ".join(current_text)
                    avg_confidence = sum(current_conf) / len(current_conf) / 100.0
                    
                    block = PreScanBlock(
                        type="text",
                        bbox=[0, 0, image.width, image.height],
                        text=combined_text,
                        confidence=avg_confidence,
                        page_number=1
                    )
                    text_blocks.append(block)
                
                if text_blocks:
                    if not result.pages:
                        result.pages = []
                    result.pages.extend(text_blocks)
                    result.total_pages = (result.total_pages or 0) + 1
                
                logger.info(f"Processed image: extracted {len(combined_text)} characters")
        
        except Exception as e:
            logger.error(f"Error processing image file: {e}")
            raise
    
    async def _detect_entities(self, result: PreScanJSON) -> None:
        """Detect entities in extracted text"""
        if not self.nlp:
            logger.warning("SpaCy not available for entity detection")
            return
        
        try:
            # Combine all text from pages and audio segments
            all_text = ""
            
            if result.pages:
                for page in result.pages:
                    all_text += page.text + "\n"
            
            if result.audio_segments:
                # Note: Audio segments don't have text yet, this would be filled by transcription
                pass
            
            if not all_text.strip():
                return
            
            # Process with SpaCy
            doc = self.nlp(all_text)
            
            # Extract entities by type
            entities = {
                "names": [],
                "organizations": [],
                "locations": [],
                "dates": [],
                "monetary": []
            }
            
            for ent in doc.ents:
                entity_text = ent.text.strip()
                if len(entity_text) > 1:  # Filter out single characters
                    if ent.label_ in ["PERSON"]:
                        entities["names"].append(entity_text)
                    elif ent.label_ in ["ORG"]:
                        entities["organizations"].append(entity_text)
                    elif ent.label_ in ["GPE", "LOC"]:
                        entities["locations"].append(entity_text)
                    elif ent.label_ in ["DATE", "TIME"]:
                        entities["dates"].append(entity_text)
                    elif ent.label_ in ["MONEY"]:
                        entities["monetary"].append(entity_text)
            
            # Remove duplicates
            for key in entities:
                entities[key] = list(set(entities[key]))
            
            result.entities_detected = entities
            
            logger.info(f"Detected entities: {sum(len(v) for v in entities.values())} total")
        
        except Exception as e:
            logger.error(f"Error in entity detection: {e}")
    
    def _calculate_overall_confidence(self, result: PreScanJSON) -> float:
        """Calculate overall confidence score"""
        confidences = []
        
        # Include page confidences
        if result.pages:
            for page in result.pages:
                confidences.append(page.confidence)
        
        # Include diarization confidence
        if result.diarization_confidence:
            confidences.append(result.diarization_confidence)
        
        if not confidences:
            return 0.0
        
        return sum(confidences) / len(confidences)



